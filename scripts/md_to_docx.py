# -*- coding: utf-8 -*-
"""Minimal Markdown → DOCX converter for the project report (python-docx)."""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def _add_runs(paragraph, text: str) -> None:
    """Support **bold** segments inside a line."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=cols)
        table.style = "Table Grid"
        for r_i, row in enumerate(table_rows):
            for c_i in range(cols):
                cell = table.rows[r_i].cells[c_i]
                cell.text = row[c_i] if c_i < len(row) else ""
                if r_i == 0:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
        doc.add_paragraph("")
        table_rows = []
        in_table = False

    def flush_code() -> None:
        nonlocal code_buf
        if not code_buf:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_buf))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        code_buf = []

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if "|" in line and line.strip().startswith("|"):
            # Skip markdown separator rows.
            if re.match(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if not in_table:
                flush_table()
                in_table = True
                table_rows = []
            table_rows.append(cells)
            i += 1
            continue
        else:
            flush_table()

        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("---"):
            doc.add_paragraph("")
        elif line.startswith("* ") or line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line[2:].strip())
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\d+\.\s", "", line).strip())
        else:
            p = doc.add_paragraph()
            _add_runs(p, line.strip())
        i += 1

    flush_table()
    flush_code()
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"wrote {docx_path}")


def main() -> int:
    p = argparse.ArgumentParser()
    p.add_argument("markdown")
    p.add_argument("-o", "--output", required=True)
    args = p.parse_args()
    convert(Path(args.markdown), Path(args.output))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
